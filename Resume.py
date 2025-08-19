# -*- coding: utf-8 -*-
# main.py

import sys
import json
import os
import subprocess
import shutil
import re # Import regex for advanced sanitization
from datetime import datetime, date
import copy # Used to safely duplicate resume data for editing

try:
    import questionary
    from rich.console import Console
    from rich.panel import Panel
    from rich.table import Table
    from rich.markdown import Markdown
    import docx
except ImportError:
    # ... (installation code) ...
    print("Required libraries not found. Attempting to install them now...")
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "rich", "questionary", "python-docx"])
        import questionary
        from rich.console import Console
        from rich.panel import Panel
        from rich.table import Table
        from rich.markdown import Markdown
        import docx
        print("Libraries installed successfully. Please restart the script.")
        sys.exit(0)
    except Exception as e:
        print(f"Failed to install required libraries: {e}")
        print("Please try installing them manually by running: pip install rich questionary python-docx")
        sys.exit(1)

try:
    import google.generativeai as genai
except ImportError:
    print("Error: google-generativeai library is missing.")
    print("Please install it by running: pip install google-generativeai")
    sys.exit(1)


# --- Configuration ---
API_KEY = os.getenv("GEMINI_API_KEY")
CONSOLE = Console()
JOB_TRACKER_FILE = "job_tracker.json"
NETWORKING_LOG_FILE = "networking_log.json"

# --- Generic Data Handling ---
def load_data(file_path):
    """Safely loads data from a JSON file."""
    if not os.path.exists(file_path):
        return []
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return json.load(f)
    except (json.JSONDecodeError, IOError):
        CONSOLE.print(f"[bold red]Warning: Could not read or parse {file_path}. Starting fresh.[/bold red]")
        return []

def save_data(data, file_path):
    """Saves data to a JSON file."""
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4, ensure_ascii=False)
    except IOError as e:
        CONSOLE.print(f"[bold red]Error saving data to {file_path}: {e}[/bold red]")


# --- Utility Functions ---
def sanitize_for_latex(text):
    """Escapes special LaTeX characters in a given string."""
    if not text: return ""
    # FIX: Corrected escaping for special characters (e.g., r'\\&' to r'\&')
    conv = {
        '&': r'\&', '%': r'\%', '$': r'\$', '#': r'\#', '_': r'\_',
        '{': r'\{', '}': r'\}', '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}', '\\': r'\textbackslash{}'
    }
    regex = re.compile('|'.join(re.escape(key) for key in sorted(conv.keys(), key=len, reverse=True)))
    return regex.sub(lambda match: conv[match.group()], text)

def get_multiline_input(prompt):
    """Gathers multi-line input from the user."""
    lines = []
    CONSOLE.print(f"[cyan]{prompt} (Enter a blank line to finish):[/cyan]")
    while True:
        line = input("> ")
        if line == "": break
        lines.append(line)
    return lines

# --- AI Integration Functions ---
def call_gemini_api(prompt):
    """Generic function to call the Gemini API."""
    if not API_KEY:
        CONSOLE.print(Panel("[bold yellow]API Key not found. Please set the GEMINI_API_KEY environment variable.[/bold yellow]", title="AI Assistant Notice", border_style="yellow"))
        return None
    with CONSOLE.status("[bold green]Connecting to AI assistant...[/bold green]"):
        try:
            genai.configure(api_key=API_KEY)
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(prompt)
            clean_response = response.text.strip().replace("```json", "").replace("```", "")
            return clean_response
        except Exception as e:
            CONSOLE.print(f"[bold red]An error occurred with the AI service: {e}[/bold red]")
            return None

def get_ai_summary(resume_data, tone):
    """Gets AI suggestions for a resume summary with a specified tone."""
    experience_text = "\n".join([f"- {entry['title']} at {entry['company']}: {' '.join(entry['accomplishments'])}" for entry in resume_data.get("experience", [])])
    prompt = (
        f"You are an expert resume writer. Based on the following experience, write a professional, "
        f"results-oriented summary of 2-3 sentences for a resume. The tone should be '{tone}'. "
        "Return ONLY the summary text.\n\n"
        "Experience:\n" + experience_text
    )
    ai_summary = call_gemini_api(prompt)
    return ai_summary.strip() if ai_summary else None

def analyze_job_description(resume_data, job_description):
    """Analyzes a resume against a job description using a structured JSON prompt."""
    resume_text = json.dumps(resume_data, indent=2)
    prompt = f"""
You are an expert career coach and ATS (Applicant Tracking System) analyst.
Analyze my resume against the provided job description and return a JSON object.
Do not include any text, notes, or markdown formatting outside of the JSON object itself.

The JSON object must have two keys:
1. "analysis": A concise, bulleted markdown string that identifies keyword gaps, suggests tone adjustments, and provides 3-5 high-level, actionable recommendations for the user to consider.
2. "suggested_edits": An array of objects, where each object represents a specific, safe change that can be made automatically. Each object must have these keys:
    - "section": The resume section to change (e.g., "summary", "experience").
    - "entry_index": For the "experience" section, the index of the job to modify (e.g., 0 for the first job). For "summary", this should be 0.
    - "field": The specific field to change (e.g., "accomplishments").
    - "original_text": An exact accomplishment string from the resume to be replaced.
    - "suggested_text": The new, keyword-optimized text. Only suggest changes that rephrase existing facts to align with the job description. Do not invent new information or skills.

RESUME DATA:
```json
{resume_text}
```

JOB DESCRIPTION:
---
{job_description}
---
"""
    return call_gemini_api(prompt)

def get_ai_cover_letter_body(resume_data, cl_info, tone):
    """Gets AI suggestions for cover letter body paragraphs with a specified tone."""
    experience_text = "\n".join([f"- {entry['title']} at {entry['company']}: {' '.join(entry['accomplishments'])}" for entry in resume_data.get("experience", [])])
    prompt = f"""
You are an expert career coach writing a compelling, personalized cover letter body.
The tone should be professional, confident, and genuinely '{tone}'.
Draft 2-3 paragraphs. Do not write the salutation (Dear...) or the closing (Sincerely,...).
Return ONLY the body paragraphs.

**Candidate Information:**
- Applying for: {cl_info['job_title']} at {cl_info['company']}
- Relevant Experience Summary: {experience_text}

**Key Personalization Points from the Candidate:**
- Most Relevant Accomplishment: {cl_info['relevant_accomplishment']}
- Excitement about the Company: {cl_info['company_excitement']}
- Specific Skill Example: {cl_info['skill_example']}

Based on all this information, write the cover letter body. Weave the personalization points into the narrative naturally.
Connect the candidate's past accomplishments to the future needs of the role.
"""
    ai_body = call_gemini_api(prompt)
    return ai_body.strip() if ai_body else None


# --- Data Gathering, Profile Management ---
def get_contact_info():
    CONSOLE.print(Panel("[bold]Contact Information[/bold]", expand=False))
    return {"full_name": questionary.text("Enter your full name:").ask(), "email": questionary.text("Enter your email address:").ask(), "phone": questionary.text("Enter your phone number:").ask(), "linkedin": questionary.text("Enter your full LinkedIn profile URL:").ask(), "photo_path": questionary.path("Enter the path to your professional photo (optional, press Enter to skip):").ask()}
def get_summary_section(resume_data):
    CONSOLE.print(Panel("[bold]Summary Section[/bold]", expand=False))
    choice = questionary.select("How would you like to create the summary?", choices=["Write my own", "Get an AI-generated suggestion"]).ask()
    if choice == "Get an AI-generated suggestion":
        tone = questionary.select("Choose a tone for the summary:", choices=["Professional", "Enthusiastic", "Direct", "Creative"]).ask()
        if not tone: return ' '.join(get_multiline_input("Enter your summary paragraph."))
        ai_summary = get_ai_summary(resume_data, tone)
        if ai_summary:
            CONSOLE.print(Panel(ai_summary, title=f"[bold cyan]AI Suggestion ({tone})[/bold cyan]", border_style="cyan"))
            if questionary.confirm("Use this summary?").ask(): return ai_summary
    return ' '.join(get_multiline_input("Enter your summary paragraph."))
def get_experience():
    CONSOLE.print(Panel("[bold]Professional Experience[/bold]", expand=False))
    entries = []
    while True:
        entries.append({"title": questionary.text("Position title:").ask(), "company": questionary.text("Company name:").ask(), "location": questionary.text("Location:").ask(), "dates": questionary.text("Employment dates:").ask(), "accomplishments": get_multiline_input("Enter your accomplishments for this role.")})
        if not questionary.confirm("Add another experience entry?").ask(): break
    return entries
def get_education():
    CONSOLE.print(Panel("[bold]Education[/bold]", expand=False))
    entries = []
    while True:
        entries.append({"institution": questionary.text("Institution name:").ask(), "degree": questionary.text("Degree:").ask(), "location": questionary.text("Location:").ask(), "dates": questionary.text("Dates (e.g., 2023 -- 2025):").ask()})
        if not questionary.confirm("Add another education entry?").ask(): break
    return entries
def get_skills():
    CONSOLE.print(Panel("[bold]Skills[/bold]", expand=False))
    if not questionary.confirm("Add a skills section?").ask(): return {}
    return {"technical": get_multiline_input("Enter your Technical Skills."), "professional": get_multiline_input("Enter your Professional Skills.")}
def get_languages():
    CONSOLE.print(Panel("[bold]Languages[/bold]", expand=False))
    if questionary.confirm("Add a languages section?").ask(): return get_multiline_input("Enter languages and proficiency (e.g., English: Native).")
    return []
def save_profile(data):
    CONSOLE.print("\n")
    filename = questionary.text("Enter filename to save profile:", default=f"{data.get('contact', {}).get('full_name', 'profile').replace(' ', '_')}.json").ask()
    if not filename: return
    filepath = os.path.join("profiles", filename)
    save_data(data, filepath)
    CONSOLE.print(f"\n[bold green]Profile saved to {filepath}[/bold green]")
def load_profile():
    profiles = [f for f in os.listdir("profiles") if f.endswith('.json')]
    if not profiles:
        CONSOLE.print("[yellow]No saved profiles found in the 'profiles' directory.[/yellow]")
        return None
    chosen_profile = questionary.select("Choose a base profile to use:", choices=profiles).ask()
    if chosen_profile:
        filepath = os.path.join("profiles", chosen_profile)
        CONSOLE.print(f"[green]Loading {chosen_profile}...[/green]")
        return load_data(filepath)
    return None

# --- Job & Networking Tracker Functions ---
def check_and_add_contact(name, company, role=""):
    """Checks if a contact exists and prompts to add them if they don't."""
    if not name or name.lower() == 'hiring team': return
    contacts = load_data(NETWORKING_LOG_FILE)
    exists = any(c.get('name', '').lower() == name.lower() and c.get('company', '').lower() == company.lower() for c in contacts)
    if not exists:
        if questionary.confirm(f"Would you like to add '{name}' from '{company}' to your networking log?").ask():
            new_contact = {"name": name, "company": company, "role": role, "date_contacted": datetime.now().strftime("%Y-%m-%d"), "notes": "Automatically added from job application/cover letter."}
            contacts.append(new_contact)
            save_data(contacts, NETWORKING_LOG_FILE)
            CONSOLE.print(f"[bold green]✔[/bold green] Contact '{name}' added to networking log.")
def add_new_application():
    """Workflow for adding a new job application to the tracker."""
    CONSOLE.print(Panel("[bold]Add New Job Application[/bold]", expand=False))
    applications = load_data(JOB_TRACKER_FILE)
    company = questionary.text("Company Name:").ask(); job_title = questionary.text("Job Title:").ask(); contact_person = questionary.text("Contact Person (optional, press Enter to skip):").ask()
    new_app = {"company": company, "job_title": job_title, "contact_person": contact_person, "date_applied": datetime.now().strftime("%Y-%m-%d"), "status": "Applied", "follow_up_date": questionary.text("Set a follow-up reminder date? (YYYY-MM-DD, optional):").ask(), "notes": ' '.join(get_multiline_input("Enter any notes (optional):")), "resume_version": ""}
    applications.append(new_app)
    save_data(applications, JOB_TRACKER_FILE)
    CONSOLE.print(f"\n[bold green]Successfully added application for {new_app['job_title']} at {new_app['company']}.[/bold green]")
    if contact_person: check_and_add_contact(contact_person, company)
def view_all_applications():
    """Displays all job applications in a formatted table."""
    CONSOLE.print(Panel("[bold]All Job Applications[/bold]", expand=False))
    applications = load_data(JOB_TRACKER_FILE)
    if not applications:
        CONSOLE.print("[yellow]No applications found in the tracker.[/yellow]"); return
    table = Table(show_header=True, header_style="bold magenta")
    table.add_column("Company", style="dim", width=20); table.add_column("Job Title", width=30); table.add_column("Contact"); table.add_column("Date Applied"); table.add_column("Status"); table.add_column("Follow-up"); table.add_column("Resume File"); table.add_column("Notes")
    for app in applications: table.add_row(app.get("company", "N/A"), app.get("job_title", "N/A"), app.get("contact_person", ""), app.get("date_applied", "N/A"), app.get("status", "N/A"), app.get("follow_up_date", ""), app.get("resume_version", ""), app.get("notes", ""))
    CONSOLE.print(table)
def add_new_contact():
    """Adds a new contact to the networking log."""
    CONSOLE.print(Panel("[bold]Add New Contact[/bold]", expand=False))
    contacts = load_data(NETWORKING_LOG_FILE)
    new_contact = {"name": questionary.text("Contact Name:").ask(), "company": questionary.text("Company:").ask(), "role": questionary.text("Role/Title:").ask(), "date_contacted": datetime.now().strftime("%Y-%m-%d"), "notes": ' '.join(get_multiline_input("Enter notes about your conversation:"))}
    contacts.append(new_contact)
    save_data(contacts, NETWORKING_LOG_FILE)
    CONSOLE.print(f"\n[bold green]Successfully added contact: {new_contact['name']}.[/bold green]")
def view_all_contacts():
    """Displays all networking contacts."""
    CONSOLE.print(Panel("[bold]Networking Log[/bold]", expand=False))
    contacts = load_data(NETWORKING_LOG_FILE)
    if not contacts:
        CONSOLE.print("[yellow]No contacts found in the log.[/yellow]"); return
    table = Table(show_header=True, header_style="bold cyan")
    table.add_column("Name", width=20); table.add_column("Company"); table.add_column("Role"); table.add_column("Date Contacted"); table.add_column("Notes", width=40)
    for contact in contacts: table.add_row(contact.get("name"), contact.get("company"), contact.get("role"), contact.get("date_contacted"), contact.get("notes"))
    CONSOLE.print(table)
def check_for_reminders():
    """Checks for and displays job application follow-up reminders."""
    applications = load_data(JOB_TRACKER_FILE)
    reminders = []
    today = date.today()
    for app in applications:
        if app.get("follow_up_date"):
            try:
                follow_up_dt = datetime.strptime(app["follow_up_date"], "%Y-%m-%d").date()
                if follow_up_dt <= today: reminders.append(f"- Follow up on [bold]{app.get('job_title')}[/bold] at [bold]{app.get('company')}[/bold]")
            except ValueError: continue
    if reminders:
        reminder_text = "\n".join(reminders)
        CONSOLE.print(Panel(reminder_text, title="[bold yellow]🔔 Follow-up Reminders[/bold yellow]", border_style="yellow"))

# --- LaTeX Section Builders & Generation ---
# Resume.py

def build_summary_section(data):
    """Builds the LaTeX code for the summary section."""
    if not data.get("summary"):
        return ""
    # FIX: Removed the non-standard \justify{...} command.
    # LaTeX justifies paragraphs by default.
    return f"\\section*{{Summary}}\n{sanitize_for_latex(data['summary'])}\\vspace{{10pt}}\n"

# Resume.py

def build_experience_section(data):
    """Builds the LaTeX code for the professional experience section."""
    if not data.get("experience"):
        return ""
    experience_str = "\\section*{Professional Experience}\n\\begin{itemize}[leftmargin=0.15in, label={}]\n"
    for entry in data["experience"]:
        title = sanitize_for_latex(entry.get('title', ''))
        dates = sanitize_for_latex(entry.get('dates', ''))
        company = sanitize_for_latex(entry.get('company', ''))
        location = sanitize_for_latex(entry.get('location', ''))
        
        # FIX: Replaced problematic '\\textbackslash{}' with a simple '/' for clarity.
        title = title.replace('\\textbackslash{}', '/')

        experience_str += f"  \\item\n    \\begin{{tabular*}}{{\\textwidth}}{{@{{\\extracolsep{{\\fill}}}}l r}}\n      \\textbf{{\\large {title}}} & {{\\small {dates}}} \\\\\n      \\textit{{\\small {company}}} & \\textit{{\\small {location}}} \\\\\n    \\end{{tabular*}}\\vspace{{-2pt}}\n"
        
        if entry.get('accomplishments'):
            experience_str += "    \\begin{itemize}[leftmargin=0.2in, topsep=0pt, itemsep=-2pt]\n"
            for acc in entry['accomplishments']:
                # FIX: Removed the non-standard \justify{...} command from here.
                # Also, replaced '\\%' which can cause issues with a simple '%' after sanitizing.
                sanitized_acc = sanitize_for_latex(acc).replace('\\textbackslash{}\\%', '\\%')
                experience_str += f"      \\item \\small{{{sanitized_acc}}}\n"
            experience_str += "    \\end{itemize}\n"
    experience_str += "\\end{itemize}\n\n"
    return experience_str

def build_education_section(data):
    if not data.get("education"): return ""
    education_str = "\\section*{Education}\n\\begin{itemize}[leftmargin=0.15in, label={}]\n"
    for entry in data["education"]:
        institution = sanitize_for_latex(entry.get('institution', '')); dates = sanitize_for_latex(entry.get('dates', '')); degree = sanitize_for_latex(entry.get('degree', '')); location = sanitize_for_latex(entry.get('location', ''))
        education_str += f"  \\item\n    \\begin{{tabular*}}{{\\textwidth}}{{@{{\\extracolsep{{\\fill}}}}l r}}\n      \\textbf{{\\large {institution}}} & {{\\small {dates}}} \\\\\n      \\textit{{\\small {degree}}} & \\textit{{\\small {location}}} \\\\\n    \\end{{tabular*}}\\vspace{{-2pt}}\n"
    education_str += "\\end{itemize}\n\n"
    return education_str
# Resume.py

def build_skills_section(data):
    """Builds the LaTeX code for the skills section."""
    skills = data.get("skills")
    if not skills or (not skills.get("technical") and not skills.get("professional")):
        return ""

    skills_str = "\\section*{Skills}\n\\noindent\n"
    tech_skills = skills.get("technical", [])
    prof_skills = skills.get("professional", [])

    def clean_skill(skill):
        """Cleans up common problematic text before sanitizing."""
        # FIX: Replace awkward line breaks and backslashes with a simple " & "
        skill = skill.replace('\\textbackslash{}\\&', ' & ')
        return sanitize_for_latex(skill)

    tech_col_width = "0.5" if tech_skills and prof_skills else "1.0"
    prof_col_width = "0.5" if tech_skills and prof_skills else "1.0"

    if tech_skills:
        skills_str += f"\\begin{{minipage}}[t]{{{tech_col_width}\\textwidth}}\n    \\begin{{itemize}}[leftmargin=0.15in, label={{}}, noitemsep, topsep=0pt]\n        \\item \\textbf{{Technical Skills}}\n        \\begin{{itemize}}[leftmargin=0.2in, topsep=2pt, itemsep=-2pt]\n"
        for skill in tech_skills:
            skills_str += f"            \\item {clean_skill(skill)}\n"
        skills_str += "        \\end{itemize}\n    \\end{itemize}\n\\end{minipage}"
        if prof_skills:
            skills_str += "%\n"

    if prof_skills:
        skills_str += f"\\begin{{minipage}}[t]{{{prof_col_width}\\textwidth}}\n    \\begin{{itemize}}[leftmargin=0.15in, label={{}}, noitemsep, topsep=0pt]\n        \\item \\textbf{{Professional Skills}}\n        \\begin{{itemize}}[leftmargin=0.2in, topsep=2pt, itemsep=-2pt]\n"
        for skill in prof_skills:
            skills_str += f"            \\item {clean_skill(skill)}\n"
        skills_str += "        \\end{itemize}\n    \\end{itemize}\n\\end{minipage}\n\n"
        
    return skills_str

def build_languages_section(data):
    if not data.get("languages"): return ""
    lang_str = "\\section*{Languages}\n\\begin{itemize}[leftmargin=0.15in, topsep=0pt]\n"
    for lang in data["languages"]: lang_str += f"    \\item {sanitize_for_latex(lang)}\n"
    lang_str += "\\end{itemize}\n"
    return lang_str
def build_photo_block(data):
    photo_path = data.get("contact", {}).get("photo_path")
    if photo_path and os.path.exists(photo_path):
        photo_filename = os.path.basename(photo_path)
        dest_path = os.path.join("output", photo_filename)
        shutil.copy(photo_path, dest_path)
        safe_path = photo_filename.replace('\\', '/')
        return f"\\includegraphics[width=0.8\\textwidth]{{{safe_path}}}"
    return ""

def compile_latex_to_pdf(latex_code, filename_base):
    if not shutil.which("pdflatex"):
        CONSOLE.print(Panel("[bold yellow]Could not find 'pdflatex'. Please install a LaTeX distribution to generate PDFs.[/bold yellow]", title="PDF Compilation Skipped"))
        return None
    output_dir = "output"; tex_filepath = os.path.join(output_dir, f"{filename_base}.tex")
    try:
        with open(tex_filepath, "w", encoding="utf-8") as f: f.write(latex_code)
    except IOError as e: CONSOLE.print(f"[bold red]Error writing .tex file: {e}[/bold red]"); return None
    with CONSOLE.status(f"[bold green]Compiling {filename_base}.pdf...[/bold green]"):
        try:
            cmd_args = ["pdflatex", "-interaction=nonstopmode", "-output-directory", output_dir, tex_filepath]
            # Run twice to ensure table of contents, etc., are correct
            subprocess.run(cmd_args, check=True, capture_output=True, text=True)
            subprocess.run(cmd_args, check=True, capture_output=True, text=True)
            pdf_path = os.path.join(output_dir, f"{filename_base}.pdf")
            CONSOLE.print(f"[bold green]✔ Successfully created {pdf_path}[/bold green]")
            return f"{filename_base}.pdf"
        except subprocess.CalledProcessError as e:
            log_path = os.path.join(output_dir, f"{filename_base}.log")
            CONSOLE.print(f"[bold red]Error during PDF compilation. Check log: {log_path}[/bold red]")
            # FIX: Print the actual captured error from pdflatex for better debugging
            if e.stderr:
                CONSOLE.print(Panel(e.stderr, title="[bold red]Compiler Error Output[/bold red]", border_style="red"))
            return None
        finally:
            # Clean up auxiliary files
            for ext in [".aux", ".log"]:
                aux_file = os.path.join(output_dir, filename_base + ext)
                if os.path.exists(aux_file):
                    try: os.remove(aux_file)
                    except OSError: pass

def generate_resume_latex(data, template_name):
    template_path = os.path.join("templates", f"{template_name.lower().replace(' ', '_')}.tex")
    if not os.path.exists(template_path):
        CONSOLE.print(f"[bold red]Error: Template file not found: {template_path}[/bold red]")
        return None
    with open(template_path, 'r', encoding='utf-8') as f: template_string = f.read()
    contact = data.get("contact", {})
    linkedin_display = contact.get('linkedin', '').replace('https://www.', '').replace('http://www.', '')
    placeholders = {"FULL_NAME": sanitize_for_latex(contact.get('full_name', '')), "EMAIL": sanitize_for_latex(contact.get('email', '')), "PHONE": sanitize_for_latex(contact.get('phone', '')), "LINKEDIN": sanitize_for_latex(contact.get('linkedin', '')), "LINKEDIN_DISPLAY": sanitize_for_latex(linkedin_display), "SUMMARY_SECTION": build_summary_section(data), "EXPERIENCE_SECTION": build_experience_section(data), "EDUCATION_SECTION": build_education_section(data), "SKILLS_SECTION": build_skills_section(data), "LANGUAGES_SECTION": build_languages_section(data), "PHOTO_BLOCK": build_photo_block(data)}
    return template_string.format(**placeholders)

def generate_cover_letter_latex(resume_data, cl_data):
    """Generates a LaTeX string for a cover letter from a template."""
    template_path = os.path.join("templates", "cover_letter.tex")
    if not os.path.exists(template_path):
        CONSOLE.print(f"[bold red]Error: Cover letter template not found at {template_path}[/bold red]")
        return None
    with open(template_path, 'r', encoding='utf-8') as f: template_string = f.read()
    
    contact = resume_data.get("contact", {})
    placeholders = {
        "FULL_NAME": sanitize_for_latex(contact.get('full_name', '')),
        "EMAIL": sanitize_for_latex(contact.get('email', '')),
        "PHONE": sanitize_for_latex(contact.get('phone', '')),
        "LINKEDIN": sanitize_for_latex(contact.get('linkedin', '')),
        "HIRING_MANAGER": sanitize_for_latex(cl_data.get('hiring_manager', 'Hiring Team')),
        "COMPANY": sanitize_for_latex(cl_data.get('company', '')),
        "COMPANY_ADDRESS": sanitize_for_latex(cl_data.get('company_address', '')),
        "JOB_TITLE": sanitize_for_latex(cl_data.get('job_title', '')),
        "BODY": cl_data.get('body', '')
    }
    return template_string.format(**placeholders)

# --- DOCX Generation ---
def generate_docx_resume(data, filename_base):
    """Generates a .docx resume from a template by building it programmatically."""
    template_path = os.path.join("templates", "resume_template.docx")
    if not os.path.exists(template_path):
        CONSOLE.print(f"[bold red]Error: Word template 'resume_template.docx' not found in 'templates' folder.[/bold red]")
        return

    doc = docx.Document(template_path)
    
    replacers = {
        "{FULL_NAME}": lambda p, d: populate_docx_full_name(p, d),
        "{CONTACT_INFO}": lambda p, d: p.clear().add_run(f"{d.get('contact', {}).get('email', '')} | {d.get('contact', {}).get('phone', '')} | {d.get('contact', {}).get('linkedin', '')}"),
        "{SUMMARY_CONTENT}": lambda p, d: p.clear().add_run(d.get('summary', '')),
        "{EXPERIENCE_CONTENT}": lambda p, d: populate_docx_experience(p, d.get("experience", [])),
        "{EDUCATION_CONTENT}": lambda p, d: populate_docx_education(p, d.get("education", [])),
        "{SKILLS_CONTENT}": lambda p, d: populate_docx_skills(p, d.get("skills", {})),
        "{LANGUAGES_CONTENT}": lambda p, d: p.clear().add_run(", ".join(d.get("languages", [])))
    }

    placeholder_paras = []
    for para in doc.paragraphs:
        for placeholder in replacers.keys():
            if placeholder in para.text:
                placeholder_paras.append((placeholder, para))
                break

    for placeholder, para in placeholder_paras:
        replacers[placeholder](para, data)

    output_path = os.path.join("output", f"{filename_base}.docx")
    try:
        doc.save(output_path)
        CONSOLE.print(f"[bold green]Successfully created {output_path}[/bold green]")
    except IOError as e:
        CONSOLE.print(f"[bold red]Error saving .docx file: {e}[/bold red]")

def populate_docx_full_name(p, data):
    """Helper function to populate the full name with bolding."""
    p.clear()
    run = p.add_run(data.get("contact", {}).get('full_name', ''))
    run.bold = True

def populate_docx_experience(p, experience_data):
    """Clears a placeholder paragraph and inserts formatted experience data before it."""
    for entry in reversed(experience_data):
        p.insert_paragraph_before("")
        for acc in reversed(entry.get('accomplishments', [])):
            p.insert_paragraph_before(f"• {acc}")
        title_paragraph = p.insert_paragraph_before(f"{entry.get('title', '')}, {entry.get('company', '')} ({entry.get('dates', '')})")
        if title_paragraph.runs:
            title_paragraph.runs[0].bold = True
    
    p_element = p._element
    p_element.getparent().remove(p_element)

def populate_docx_education(p, education_data):
    """Clears a placeholder and populates it with formatted education data."""
    p.clear()
    edu_text = "\n".join([f"{entry.get('degree', '')}, {entry.get('institution', '')} ({entry.get('dates', '')})" for entry in education_data])
    p.add_run(edu_text)

def populate_docx_skills(p, skills_data):
    """Clears a placeholder and populates it with formatted skills data."""
    p.clear()
    if skills_data.get("technical"):
        p.add_run("Technical: ").bold = True
        p.add_run(", ".join(skills_data["technical"]) + "\n")
    if skills_data.get("professional"):
        p.add_run("Professional: ").bold = True
        p.add_run(", ".join(skills_data["professional"]))


# --- Main Application Logic & Workflows ---

def create_new_resume():
    """Workflow for creating a new resume from scratch."""
    resume_data = {}
    resume_data["contact"] = get_contact_info()
    resume_data["experience"] = get_experience()
    resume_data["summary"] = get_summary_section(resume_data)
    resume_data["education"] = get_education()
    resume_data["skills"] = get_skills()
    resume_data["languages"] = get_languages()
    
    if questionary.confirm("Do you want to save this profile for future use?").ask(): 
        save_profile(resume_data)
    return resume_data

def apply_ai_edits(resume_data, edits):
    """Safely applies a list of AI-suggested edits to the resume data."""
    editable_resume = copy.deepcopy(resume_data)
    for edit in edits:
        try:
            section = edit['section']
            if section == 'summary':
                if editable_resume.get('summary', '') == edit['original_text']:
                    editable_resume['summary'] = edit['suggested_text']
                    CONSOLE.print(f"[green]✔[/green] Updated summary.")
            elif section == 'experience':
                accomplishments = editable_resume['experience'][edit['entry_index']]['accomplishments']
                if edit['original_text'] in accomplishments:
                    accomplishments[accomplishments.index(edit['original_text'])] = edit['suggested_text']
                    CONSOLE.print(f"[green]✔[/green] Updated accomplishment in job #{edit['entry_index'] + 1}.")
        except (KeyError, IndexError): CONSOLE.print(f"[yellow]Could not apply an edit due to data mismatch: {edit}[/yellow]"); continue
    return editable_resume

def generation_workflow(resume_data, filename_base=None):
    """Handles the choice of generation format (PDF, DOCX, or both)."""
    if not resume_data: return
    
    if not filename_base:
        # FIX: Use .get() for safer dictionary access to prevent errors if resume_data is not a valid dict.
        full_name = resume_data['contact']['full_name'] if isinstance(resume_data, dict) and 'contact' in resume_data and 'full_name' in resume_data['contact'] else 'user'
        filename_base = f"resume_{full_name.replace(' ', '_')}"
    
    format_choices = questionary.checkbox("Select output format(s):", choices=["PDF (via LaTeX)", "Word (DOCX)"]).ask()
    if not format_choices: return

    if "PDF (via LaTeX)" in format_choices:
        template_choice = questionary.select("Choose a LaTeX template:", choices=["Modern", "Classic", "Photo Professional", "Creative"]).ask()
        if template_choice:
            latex_code = generate_resume_latex(resume_data, template_choice)
            if latex_code:
                compile_latex_to_pdf(latex_code, filename_base)

    if "Word (DOCX)" in format_choices:
        generate_docx_resume(resume_data, filename_base)


def job_description_workflow(application):
    """Workflow for analyzing a job description for a specific application."""
    CONSOLE.print(Panel(f"Analyzing for: [bold]{application['job_title']} at {application['company']}[/bold]", expand=False, border_style="cyan"))
    CONSOLE.print("[cyan]First, load the base resume profile you want to tailor.[/cyan]")
    resume_data = load_profile()
    if not resume_data: return

    job_desc = ' '.join(get_multiline_input("Paste the job description below."))
    if not job_desc: return

    ai_response_str = analyze_job_description(resume_data, job_desc)
    if not ai_response_str:
        CONSOLE.print("[bold red]Failed to get a response from the AI.[/bold red]"); return
    try:
        ai_data = json.loads(ai_response_str)
        analysis_text = ai_data.get("analysis", "No analysis provided.")
        suggested_edits = ai_data.get("suggested_edits", [])
        CONSOLE.print(Panel(Markdown(analysis_text), title="[bold cyan]AI Analysis[/bold cyan]", border_style="cyan"))
        
        if suggested_edits and questionary.confirm("The AI has suggested automated edits. Apply them and generate a new resume?").ask():
            updated_resume_data = apply_ai_edits(resume_data, suggested_edits)
            # FIX: Use .get() for safer dictionary access.
            full_name = updated_resume_data.get('contact', {}).get('full_name', 'user')
            company_name = application.get('company', 'company')
            filename_base = f"resume_{full_name.replace(' ', '_')}_for_{company_name.replace(' ', '_')}"
            generation_workflow(updated_resume_data, filename_base)
            # Here you could update the job tracker with the new filename, but batch workflow is better for that.
            
    except json.JSONDecodeError:
        CONSOLE.print("[bold red]Could not parse the AI's response. The raw response was:[/bold red]")
        CONSOLE.print(ai_response_str)

def cover_letter_workflow(application):
    """Workflow for generating a personalized cover letter for a specific application."""
    CONSOLE.print(Panel(f"Generating Cover Letter for: [bold]{application['job_title']} at {application['company']}[/bold]", expand=False, border_style="green"))
    CONSOLE.print("[cyan]First, load a resume profile to get your contact info and experience.[/cyan]")
    resume_data = load_profile()
    if not resume_data: return

    # Pre-fill data from the application tracker
    cl_info = {
        "hiring_manager": application.get("contact_person") or questionary.text("Hiring Manager's Name (or 'Hiring Team'):").ask(),
        "company": application.get("company"),
        "job_title": application.get("job_title"),
        "company_address": ' '.join(get_multiline_input("Company Address:")),
    }

    check_and_add_contact(cl_info["hiring_manager"], cl_info["company"])
    tone = questionary.select("Choose a tone for the cover letter:", choices=["Professional", "Enthusiastic", "Direct", "Creative"]).ask()
    if not tone: return

    CONSOLE.print("\n[cyan]To make the letter personal, please answer a few more questions:[/cyan]")
    cl_info["relevant_accomplishment"] = questionary.text("Which accomplishment from your resume is most relevant to this job's top requirement?").ask()
    cl_info["company_excitement"] = questionary.text("What is one specific thing about this company's mission or product that excites you?").ask()
    cl_info["skill_example"] = questionary.text("Briefly describe a time you used a key skill mentioned in the job description (e.g., 'project management').").ask()

    ai_body = get_ai_cover_letter_body(resume_data, cl_info, tone)
    if ai_body:
        cl_info['body'] = sanitize_for_latex(ai_body)
        CONSOLE.print(Panel(ai_body, title=f"[bold green]AI-Generated Cover Letter Body ({tone})[/bold green]", border_style="green"))
        if questionary.confirm("Do you want to use this draft and generate the full letter?").ask():
            latex_code = generate_cover_letter_latex(resume_data, cl_info)
            if latex_code:
                # FIX: Use .get() for safer dictionary access.
                full_name = resume_data['contact']['full_name'] if isinstance(resume_data, dict) and 'contact' in resume_data and 'full_name' in resume_data['contact'] else 'user'
                company_name = cl_info.get('company', 'company')
                filename_base = f"Cover_Letter_{full_name.replace(' ', '_')}_for_{company_name.replace(' ', '_')}"
                compile_latex_to_pdf(latex_code, filename_base)
    else:
        CONSOLE.print("[bold red]Could not generate the cover letter body.[/bold red]")

def batch_resume_workflow():
    """Workflow for tailoring multiple resumes in a single batch."""
    # ... (This function remains unchanged) ...
    CONSOLE.print(Panel("[bold]Batch Resume Tailoring[/bold]", expand=False, border_style="yellow"))
    CONSOLE.print("[cyan]First, load the base resume profile you want to tailor.[/cyan]")
    base_resume_data = load_profile()
    if not base_resume_data: return

    applications = load_data(JOB_TRACKER_FILE)
    unprocessed_apps = [app for app in applications if app.get("status") == "Applied" and not app.get("resume_version")]
    if not unprocessed_apps:
        CONSOLE.print("[yellow]No applications are ready for batch processing (must have status 'Applied' and no existing resume version).[/yellow]"); return

    app_choices = [f"{app.get('job_title')} at {app.get('company')}" for app in unprocessed_apps]
    selected_apps_str = questionary.checkbox("Select jobs to process:", choices=app_choices).ask()
    if not selected_apps_str: return

    for app_str in selected_apps_str:
        app_index = app_choices.index(app_str)
        target_app = unprocessed_apps[app_index]
        
        CONSOLE.print(Panel(f"Processing: [bold]{app_str}[/bold]", border_style="green"))
        job_desc = ' '.join(get_multiline_input(f"Paste the job description for '{app_str}':"))
        if not job_desc:
            CONSOLE.print("[yellow]Skipping due to no job description.[/yellow]"); continue

        ai_response_str = analyze_job_description(base_resume_data, job_desc)
        if not ai_response_str:
            CONSOLE.print("[red]Skipping due to AI error.[/red]"); continue
        
        try:
            ai_data = json.loads(ai_response_str)
            suggested_edits = ai_data.get("suggested_edits", [])
            if suggested_edits:
                updated_resume_data = apply_ai_edits(base_resume_data, suggested_edits)
                # FIX: Use .get() for safer dictionary access.
                full_name = base_resume_data['contact']['full_name'] if isinstance(base_resume_data, dict) and 'contact' in base_resume_data and 'full_name' in base_resume_data['contact'] else 'user'
                company_name = target_app.get('company', 'company')
                filename_base = f"resume_{full_name.replace(' ', '_')}_for_{company_name.replace(' ', '_')}"
                
                latex_code = generate_resume_latex(updated_resume_data, "Modern")
                if latex_code:
                    generated_filename = compile_latex_to_pdf(latex_code, filename_base)
                    if generated_filename:
                        for original_app in applications:
                            if original_app['company'] == target_app['company'] and original_app['job_title'] == target_app['job_title']:
                                original_app['resume_version'] = generated_filename
                                break
                        save_data(applications, JOB_TRACKER_FILE)
                        CONSOLE.print(f"[green]Updated job tracker for {app_str} with new resume file.[/green]")

            else:
                CONSOLE.print("[yellow]AI suggested no edits for this job.[/yellow]")

        except json.JSONDecodeError:
            CONSOLE.print(f"[red]Could not parse AI response for {app_str}. Skipping.[/red]")
            continue

    CONSOLE.print("\n[bold green]Batch processing complete![/bold green]")

def job_tracker_menu(application):
    """Shows the action menu for a single selected job application."""
    while True:
        CONSOLE.print(Panel(f"Selected Application: [bold]{application['job_title']} at {application['company']}[/bold]", border_style="purple"))
        choice = questionary.select(
            "What would you like to do with this application?",
            choices=[
                "Tailor Resume for this Job",
                "Generate Cover Letter for this Job",
                "View/Edit Contact Person",
                "Update Status",
                "Back to Job List"
            ]).ask()
        
        if not choice or choice == "Back to Job List": break

        if choice == "Tailor Resume for this Job":
            job_description_workflow(application)
        elif choice == "Generate Cover Letter for this Job":
            cover_letter_workflow(application)
        elif choice == "View/Edit Contact Person":
            # This is a placeholder for a more advanced edit feature
            CONSOLE.print(f"Current contact: {application.get('contact_person', 'N/A')}")
            questionary.press_any_key_to_continue().ask()
        elif choice == "Update Status":
            # We need to find the application in the main list to update it
            applications = load_data(JOB_TRACKER_FILE)
            for i, app in enumerate(applications):
                if app['company'] == application['company'] and app['job_title'] == application['job_title']:
                    new_status = questionary.select("Select new status:", choices=["Applied", "Interviewing", "Offer Received", "Rejected", "Closed"]).ask()
                    if new_status:
                        applications[i]['status'] = new_status
                        save_data(applications, JOB_TRACKER_FILE)
                        CONSOLE.print("[green]Status updated.[/green]")
                    break

def select_application_workflow():
    """The main entry point for the job tracker, allowing users to select a job to manage."""
    while True:
        CONSOLE.print(Panel("[bold]Job Application Manager[/bold]", expand=False, border_style="purple"))
        applications = load_data(JOB_TRACKER_FILE)
        if not applications:
            CONSOLE.print("[yellow]No applications found. Add one to get started.[/yellow]")
            if questionary.confirm("Add a new application now?").ask():
                add_new_application()
            return

        app_choices = [f"{app.get('job_title')} at {app.get('company')}" for app in applications]
        
        # FIX: Combine lists using '+' to avoid the .extend type error.
        choices = app_choices + [
            questionary.Separator(), 
            "View Full Table", 
            "Add New Application", 
            "Batch Tailor Resumes", 
            "Back to Main Menu"
        ]
        
        choice = questionary.select("Select an application to manage, or choose an option:", choices=choices).ask()

        if not choice or choice == "Back to Main Menu": break
        elif choice == "View Full Table": view_all_applications()
        elif choice == "Add New Application": add_new_application()
        elif choice == "Batch Tailor Resumes": batch_resume_workflow()
        else:
            # User selected a specific job
            app_index = app_choices.index(choice) # Note: Searching in app_choices, not the combined list
            selected_app = applications[app_index]
            job_tracker_menu(selected_app) # Open the action menu for that job

def main():
    """Main function to run the resume builder application."""
    check_for_reminders()
    while True:
        CONSOLE.print(Panel("[bold]AI-Powered Resume & Career Toolkit[/bold]", title="[bold]Main Menu[/bold]", border_style="blue"))
        choice = questionary.select(
            "What would you like to do?",
            choices=[
                "Manage Job Applications", # <-- New central hub
                "Create a New Base Resume Profile",
                "Load and Generate from Base Profile",
                "Networking Log",
                "Exit"
            ]).ask()
        if not choice or choice == 'Exit':
            CONSOLE.print("[bold cyan]Goodbye![/bold cyan]"); break
        
        if choice == 'Manage Job Applications':
            select_application_workflow()

        elif choice == 'Create a New Base Resume Profile':
            resume_data = create_new_resume()
            if resume_data and questionary.confirm("Generate a document from this new resume now?").ask():
                generation_workflow(resume_data)

        elif choice == 'Load and Generate from Base Profile':
            resume_data = load_profile()
            if resume_data:
                generation_workflow(resume_data)

        elif choice == 'Networking Log':
            # Simplified direct access
            view_all_contacts()
            if questionary.confirm("Add a new contact?").ask():
                add_new_contact()
        
        questionary.press_any_key_to_continue().ask()


if __name__ == "__main__":
    os.makedirs("templates", exist_ok=True); os.makedirs("profiles", exist_ok=True); os.makedirs("output", exist_ok=True)
    if not os.path.exists(JOB_TRACKER_FILE):
        with open(JOB_TRACKER_FILE, 'w') as f: json.dump([], f)
    if not os.path.exists(NETWORKING_LOG_FILE):
        with open(NETWORKING_LOG_FILE, 'w') as f: json.dump([], f)
    main()
